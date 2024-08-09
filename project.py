import os
import re
from transformers import pipeline, AutoTokenizer, AutoModelForQuestionAnswering
from docx import Document
from fpdf import FPDF
from PyPDF2 import PdfReader

# ANSI escape codes for color coding
class Colors:
    HEADER = '\033[95m'
    OKBLUE = '\033[94m'
    OKGREEN = '\033[92m'
    WARNING = '\033[93m'
    FAIL = '\033[91m'
    ENDC = '\033[0m'
    BOLD = '\033[1m'
    UNDERLINE = '\033[4m'

# Load models
summarizer = pipeline("summarization", model="sshleifer/distilbart-cnn-12-6")
qa_model_name = "deepset/roberta-base-squad2"
qa_tokenizer = AutoTokenizer.from_pretrained(qa_model_name)
qa_model = AutoModelForQuestionAnswering.from_pretrained(qa_model_name)
qa_pipeline = pipeline("question-answering", model=qa_model, tokenizer=qa_tokenizer)

# Extract text from PDF
def extract_text_from_pdf(pdf_path):
    text = ''
    with open(pdf_path, 'rb') as file:
        reader = PdfReader(file)
        for page in reader.pages:
            text += page.extract_text() + '\n\n'
    return text

# Extract text from DOCX
def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = '\n\n'.join(paragraph.text for paragraph in doc.paragraphs)
    return text

# Preprocess text
def preprocess_text(text):
    return text.strip()

# Summarize text with professionalism
def summarize_text(text, max_length=150, min_length=30):
    # Split text into chunks if too long
    max_chunk_size = 1000 * 5  # Roughly equivalent to 1000 words per chunk
    text_chunks = [text[i:i+max_chunk_size] for i in range(0, len(text), max_chunk_size)]
    
    summaries = []
    for chunk in text_chunks:
        try:
            summary = summarizer(chunk, max_length=max_length, min_length=min_length)
            summaries.append(summary[0]['summary_text'])
        except Exception as e:
            summaries.append(f"{Colors.FAIL}Error generating summary for a chunk: {e}{Colors.ENDC}")
    
    full_summary = " ".join(summaries)
    return full_summary

# Generate summary in bullet points if the document is long
def summarize_long_document(text):
    max_length = 150
    min_length = 30
    chunk_size = 1000 * 5  # Roughly 1000 words per chunk
    
    # Check if the document is long
    if len(text.split()) > 1000:  # Rough estimate for long documents
        # Generate summary for each chunk
        text_chunks = [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]
        summaries = []
        for chunk in text_chunks:
            try:
                summary = summarizer(chunk, max_length=max_length, min_length=min_length)
                summaries.append(summary[0]['summary_text'])
            except Exception as e:
                summaries.append(f"{Colors.FAIL}Error generating summary for a chunk: {e}{Colors.ENDC}")

        # Convert the summary to bullet points
        bullet_point_summary = "\n".join(f"• {s}" for s in summaries)
        return f"{Colors.OKBLUE}Summary in bullet points:{Colors.ENDC}\n{bullet_point_summary}"
    else:
        return summarize_text(text, max_length=max_length, min_length=min_length)

# Extract paragraphs related to a specific topic
def find_paragraphs_related_to_topic(text, topic):
    paragraphs = text.split('\n\n')
    related_paragraphs = [p for p in paragraphs if topic.lower() in p.lower()]
    return '\n\n'.join(related_paragraphs)

# Answer questions based on the text
def answer_question(question, text):
    answer = qa_pipeline(question=question, context=text)
    return answer['answer']

# Handle user queries
def handle_query(query, text):
    query = query.lower()
    if 'summary' in query:
        summary = summarize_long_document(text)
        return f"{Colors.OKBLUE}Summary:{Colors.ENDC}\n{summary}\n"
    elif 'first paragraph' in query:
        paragraphs = text.split('\n\n')
        if paragraphs:
            return f"{Colors.OKGREEN}First paragraph:{Colors.ENDC}\n{paragraphs[0]}\n"
        else:
            return f"{Colors.WARNING}No paragraphs found.{Colors.ENDC}\n"
    elif 'whole content' in query or 'full content' in query:
        return f"{Colors.OKGREEN}Whole content:{Colors.ENDC}\n{text}\n"
    elif 'topic' in query:
        topic = query.split('topic')[-1].strip()
        related_text = find_paragraphs_related_to_topic(text, topic)
        if related_text:
            return f"{Colors.OKGREEN}Content related to '{topic}':{Colors.ENDC}\n{related_text}\n"
        else:
            return f"{Colors.FAIL}No content related to '{topic}' found.{Colors.ENDC}\n"
    else:
        answer = answer_question(query, text)
        if answer:
            return f"{Colors.OKGREEN}Answer to your question:{Colors.ENDC}\n{answer}\n"
        else:
            return f"{Colors.FAIL}I couldn't understand your query. Please try asking something else.{Colors.ENDC}\n"

# Save output
def save_output(text, file_name, file_format):
    with open(file_name, 'w') as file:
        if file_format == 'txt':
            file.write(text)
        elif file_format == 'docx':
            doc = Document()
            doc.add_paragraph(text)
            doc.save(file_name)
        elif file_format == 'pdf':
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, text)
            pdf.output(file_name)

# Ask user to save output
def ask_to_save(text, prompt):
    save = input(f"{Colors.WARNING}{prompt} (yes/no): {Colors.ENDC}").strip().lower()
    if save == 'yes':
        file_format = input(f"{Colors.WARNING}Enter the format to save the output (txt/docx/pdf): {Colors.ENDC}").strip().lower()
        while file_format not in ['txt', 'docx', 'pdf']:
            print(f"{Colors.FAIL}Invalid format. Please enter 'txt', 'docx', or 'pdf'.{Colors.ENDC}")
            file_format = input(f"{Colors.WARNING}Enter the format to save the output (txt/docx/pdf): {Colors.ENDC}").strip().lower()
        file_dir = input(f"{Colors.WARNING}Enter the directory to save the output file: {Colors.ENDC}").strip()
        if not os.path.exists(file_dir):
            os.makedirs(file_dir)
        file_name = os.path.join(file_dir, f"output.{file_format}")
        save_output(text, file_name, file_format)
        print(f"{Colors.OKGREEN}Output saved successfully.{Colors.ENDC}\n")
    else:
        print(f"{Colors.OKBLUE}Output not saved.{Colors.ENDC}\n")

# Get input file from user
def get_input_file():
    """Prompt the user to enter the path of the PDF or DOCX file."""
    while True:
        file_path = input(f"{Colors.WARNING}Enter the path of the PDF or DOCX file: {Colors.ENDC}").strip()
        if os.path.isfile(file_path) and (file_path.endswith('.pdf') or file_path.endswith('.docx')):
            return file_path
        else:
            print(f"{Colors.FAIL}Invalid file path or file type. Please enter a valid path for a PDF or DOCX file.{Colors.ENDC}\n")

# Main function
def main():
    file_path = get_input_file()
    
    if file_path.endswith('.pdf'):
        extracted_text = extract_text_from_pdf(file_path)
    elif file_path.endswith('.docx'):
        extracted_text = extract_text_from_docx(file_path)
    
    cleaned_text = preprocess_text(extracted_text)
    
    while True:
        query = input(f"{Colors.WARNING}Enter your query (or 'exit' to quit): {Colors.ENDC}").strip()
        if query.lower() == 'exit':
            break
        
        response = handle_query(query, cleaned_text)
        print(response)
        
        if 'summary' in query or 'whole content' in query:
            ask_to_save(response, "Do you want to save the output")

if __name__ == "__main__":
    main()
